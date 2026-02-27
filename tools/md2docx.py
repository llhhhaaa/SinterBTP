# -*- coding: utf-8 -*-
"""Markdown to DOCX converter with image, table support."""
import re, os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from urllib.parse import unquote

def clean_md(text):
    """Remove markdown formatting."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    text = text.replace('\\times', 'x').replace('\\text{', '').replace('}', '')
    text = text.replace('\\frac', 'frac').replace('\\sum', 'sum')
    text = text.replace('\\mathbf{1}', '1').replace('\\leq', '<=')
    text = text.replace('\\ll', '<<').replace('\\alpha', 'alpha')
    return text.strip()

def add_runs(paragraph, text):
    """Add text with bold/italic support."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle inline $...$
            subparts = re.split(r'(\$[^$]+\$)', part)
            for sp in subparts:
                if sp.startswith('$') and sp.endswith('$'):
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sp)

def md_to_docx(md_path, docx_path):
    md_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    
    i = 0
    in_math = False
    math_buf = []
    
    while i < len(lines):
        line = lines[i]
        
        # Math block
        if line.strip().startswith('$$') and not in_math:
            if line.strip().endswith('$$') and len(line.strip()) > 4:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line.strip()[2:-2].strip())
                run.italic = True
                i += 1; continue
            in_math = True; math_buf = []; i += 1; continue
        if in_math:
            if line.strip() == '$$':
                in_math = False
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(' '.join(math_buf))
                run.italic = True
                i += 1; continue
            math_buf.append(line.strip())
            i += 1; continue
        
        # Empty
        if not line.strip():
            i += 1; continue
        
        # Heading
        hm = re.match(r'^(#{1,4})\s+(.+)', line)
        if hm:
            level = len(hm.group(1))
            doc.add_heading(hm.group(2).strip(), level=min(level, 4))
            i += 1; continue
        
        # Image - 使用更健壮的正则处理路径中的括号
        img = re.match(r'!\[([^\]]*)\]\((.+)\)$', line.strip())
        if img:
            img_path = unquote(img.group(2).strip())
            full_path = os.path.join(md_dir, img_path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(full_path):
                try:
                    run = p.add_run()
                    run.add_picture(full_path, width=Inches(5.5))
                except Exception as e:
                    p.add_run(f'[Image error: {e}]')
            else:
                run = p.add_run(f'[Image not found: {img_path}]')
                run.font.color.rgb = RGBColor(255, 0, 0)
            i += 1; continue
        
        # Table
        if '|' in line and i + 1 < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[i+1]):
            tlines = []
            while i < len(lines) and '|' in lines[i]:
                tlines.append(lines[i]); i += 1
            if len(tlines) >= 3:
                headers = [c.strip() for c in tlines[0].split('|') if c.strip()]
                rows = []
                for tl in tlines[2:]:
                    cells = [c.strip() for c in tl.split('|') if c.strip()]
                    if cells: rows.append(cells)
                ncols = len(headers)
                table = doc.add_table(rows=1+len(rows), cols=ncols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(clean_md(h))
                    run.bold = True; run.font.size = Pt(9)
                for ri, row in enumerate(rows):
                    for j in range(min(len(row), ncols)):
                        cell = table.rows[ri+1].cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(clean_md(row[j]))
                        run.font.size = Pt(9)
                        if '**' in row[j]: run.bold = True
                doc.add_paragraph()
            continue
        
        # Bullet
        if line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, line.strip()[2:])
            i += 1; continue
        
        # Numbered list
        nm = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if nm:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, nm.group(2))
            i += 1; continue
        
        # Figure caption (bold paragraph starting with **Figure or **Table)
        if line.strip().startswith('**Figure') or line.strip().startswith('**Table'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, line.strip())
            p.runs[0].font.size = Pt(10) if p.runs else None
            i += 1; continue
        
        # Note line
        if line.strip().startswith('注：') or line.strip().startswith('Note:'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip())
            run.font.size = Pt(9)
            run.italic = True
            i += 1; continue
        
        # Normal paragraph
        p = doc.add_paragraph()
        add_runs(p, line.strip())
        i += 1
    
    doc.save(docx_path)
    print(f'Saved: {docx_path}')

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python script.py input.md output.docx')
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
